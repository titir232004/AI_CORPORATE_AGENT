from docx import Document
from docx.shared import RGBColor
from typing import List, Dict
import os

def insert_comments(input_path, issues: List[Dict], output_path):
    if hasattr(input_path, "read"):
        from tempfile import NamedTemporaryFile
        tmp = NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(input_path.read())
        tmp.flush()
        tmp.close()
        input_path = tmp.name

    doc = Document(input_path)
    inserted_any = False
    for issue in issues:
        section = issue.get("section", "").lower()
        suggestion = issue.get("suggestion", "") or issue.get("issue", "")
        added = False
        if section:
            for para in doc.paragraphs:
                if section in para.text.lower():
                    run = para.add_run(f"  [REVIEW NOTE: {suggestion}]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
                    inserted_any = True
                    added = True
                    break
        if not added:
            p = doc.add_paragraph()
            run = p.add_run(f"[REVIEW NOTE - {issue.get('document','doc')}]: {suggestion}")
            run.italic = True
            run.font.color.rgb = RGBColor(0xA0, 0x00, 0x00)
            inserted_any = True
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
